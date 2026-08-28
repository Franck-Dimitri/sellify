import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_corporate_docx():
    doc = Document()

    # Page Margins (Normal: 1 inch = 1440 dxa)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("SELLIFY.ME  |  DOSSIER STRATÉGIQUE & TECHNIQUE OFFICIEL  ·  CONFIDENTIEL")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = RGBColor(100, 116, 139)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Sellify.me © 2026 — Plateforme E-Commerce de Confiance pour l'Afrique — Tous droits réservés")
        frun.font.name = "Arial"
        frun.font.size = Pt(8)
        frun.font.color.rgb = RGBColor(148, 163, 184)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # PAGE 1 : PREMIÈRE DE COUVERTURE CORPORATE
    # =========================================================================
    
    # Title Tag Box (Table)
    tbl_tag = doc.add_table(rows=1, cols=1)
    tbl_tag.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl_tag.cell(0, 0)
    set_cell_background(c, "0F172A") # Dark Navy
    set_cell_margins(c, top=200, bottom=200, left=300, right=300)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOCUMENT CORPORATE OFFICIEL DE RÉFÉRENCE")
    r.font.name = "Arial"
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(234, 179, 8) # Gold Yellow

    doc.add_paragraph().paragraph_format.space_before = Pt(30)

    # Main Project Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = p_title.add_run("SELLIFY.ME\n")
    r_main.font.name = "Arial"
    r_main.font.bold = True
    r_main.font.size = Pt(36)
    r_main.font.color.rgb = RGBColor(15, 23, 42) # Slate 900

    r_sub = p_title.add_run("La Plateforme E-Commerce de Confiance pour l'Afrique Numérique\n")
    r_sub.font.name = "Arial"
    r_sub.font.bold = True
    r_sub.font.size = Pt(16)
    r_sub.font.color.rgb = RGBColor(202, 138, 4) # Gold 600

    r_desc = p_title.add_run("Marketplace Multi-Acteurs · Séquestre Escrow Mobile Money · Logistique IA & Télémétrie · Smart-Links")
    r_desc.font.name = "Arial"
    r_desc.font.size = Pt(10.5)
    r_desc.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_before = Pt(40)

    # Metadata Box Table
    tbl_meta = doc.add_table(rows=6, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_meta.autofit = False
    
    meta_rows = [
        ("Nom du Projet :", "Sellify.me — Écosystème E-Commerce & Marketplace de Confiance"),
        ("Type de Document :", "Dossier de Présentation de Projet & Rapport d'Ingénierie Logicielle"),
        ("Version :", "Version 1.2 — Conforme au Cahier des Charges & Code Source"),
        ("Stack Technologique :", "Laravel 11 · React 19 · Inertia.js · PostgreSQL 16 · OSRM · Sellify AI 1.2 Flash"),
        ("Périmètre de Déploiement :", "Phase 1 : Cameroun (Douala & Yaoundé) | Phase 2 : Zone CEMAC & CEDEAO"),
        ("Statut & Validation :", "MVP & V1 Validés — 100% Tests Automatisés Réussis (47/47 tests)"),
    ]

    for idx, (label, val) in enumerate(meta_rows):
        c0 = tbl_meta.cell(idx, 0)
        c1 = tbl_meta.cell(idx, 1)
        c0.width = Inches(2.3)
        c1.width = Inches(4.4)

        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=120, right=120)
        set_cell_margins(c1, top=80, bottom=80, left=120, right=120)

        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor(15, 23, 42)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(val)
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_before = Pt(50)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("Direction Générale Sellify · Douala · Yaoundé · Août 2026\nDocument Strictement Confidentiel — Diffusion Restreinte")
    r_f.font.italic = True
    r_f.font.size = Pt(8.5)
    r_f.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_page_break()

    # =========================================================================
    # PAGE 2 : SOMMAIRE & EXECUTIVE SUMMARY
    # =========================================================================
    
    # Heading Helper
    def add_sec_title(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title_text)
        r.font.name = "Arial"
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(15, 23, 42)
        
        # Subtitle bottom line effect
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(6)
        r_l = p_line.add_run("━" * 55)
        r_l.font.size = Pt(8)
        r_l.font.color.rgb = RGBColor(234, 179, 8)

    def add_subsec_title(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title_text)
        r.font.name = "Arial"
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(202, 138, 4)

    def add_callout(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.cell(0, 0)
        set_cell_background(c, "FEFCE8") # Light yellow
        set_cell_margins(c, top=120, bottom=120, left=180, right=180)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(f"📌 {title.upper()}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = RGBColor(161, 98, 7)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9)
        r_b.font.color.rgb = RGBColor(66, 32, 6)
        doc.add_paragraph().paragraph_format.space_before = Pt(4)

    add_sec_title("1. Sommaire Exécutif & Vision Stratégique")

    doc.add_paragraph(
        "Sellify.me est une infrastructure technologique de confiance conçue pour accélérer la transition du commerce "
        "informel vers une économie numérique structurée en Afrique subsaharienne. En s'appuyant sur l'omniprésence du "
        "Mobile Money (Orange Money, MTN MoMo), Sellify.me introduit un compte séquestre numérique (Escrow) inviolable "
        "qui protège simultanément l'acheteur et le commerçant."
    )

    doc.add_paragraph(
        "La plateforme intègre une suite logicielle SaaS marchande permettant de gérer plusieurs boutiques, des catalogues "
        "avec variantes de stocks, des tunnels de commande 1-clic pour le Social Commerce (Smart-Links WhatsApp, Facebook, TikTok), "
        "une application livreur avec guidage GPS OSRM temps réel et un copilote d'Intelligence Artificielle universel (Sellify AI 1.2 Flash)."
    )

    add_callout(
        "Le Défi Africain Résolu",
        "En Afrique subsaharienne francophone, plus de 80% des règlements numériques sont opérés via Mobile Money, mais le e-commerce souffre "
        "d'un taux de refus à la livraison de 35% à cause du Cash on Delivery. Sellify.me sécurise 100% des fonds avant le départ du colis."
    )

    add_sec_title("2. Analyse Stratégique du Marché & Problématiques")

    doc.add_paragraph(
        "Le commerce en ligne informel en Afrique est confronté à trois goulots d'étranglement majeurs :"
    )

    # Comparison Table
    tbl_comp = doc.add_table(rows=6, cols=3)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.autofit = False

    comp_headers = ["Dimension Métier", "Commerce Informel Classique", "Solution Intégrée Sellify.me"]
    for i, h in enumerate(comp_headers):
        c = tbl_comp.cell(0, i)
        set_cell_background(c, "0F172A")
        set_cell_margins(c, top=100, bottom=100, left=100, right=100)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ("Garantie de Paiement", "Inexistante (risque d'arnaque de 40%)", "100% Séquestre Escrow Mobile Money"),
        ("Taux d'Annulation", "Supérieur à 35% (Pertes sur Cash on Delivery)", "Inférieur à 2% (Fonds consignés d'avance)"),
        ("Adressage & Itinéraire", "Appels téléphoniques à répétition", "Guidage GPS OSRM & Télémétrie temps réel"),
        ("Validation Livraison", "Verbale / Soumise à contestation", "Double validation : Code OTP Secret + Signature"),
        ("Outils de Vente", "Catalogues photos éparpillés sur WhatsApp", "Smart-Links 1-clic & Boutiques SaaS dédiées"),
    ]

    for row_idx, (d, b, a) in enumerate(comp_data, start=1):
        c0, c1, c2 = tbl_comp.cell(row_idx, 0), tbl_comp.cell(row_idx, 1), tbl_comp.cell(row_idx, 2)
        c0.width, c1.width, c2.width = Inches(2.0), Inches(2.4), Inches(2.4)
        set_cell_background(c0, "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
        set_cell_background(c1, "FFF1F2")
        set_cell_background(c2, "ECFDF5")
        set_cell_margins(c0, top=70, bottom=70, left=90, right=90)
        set_cell_margins(c1, top=70, bottom=70, left=90, right=90)
        set_cell_margins(c2, top=70, bottom=70, left=90, right=90)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(d)
        r0.font.bold = True
        r0.font.size = Pt(8)
        r0.font.color.rgb = RGBColor(15, 23, 42)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(b)
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor(190, 18, 60)

        p2 = c2.paragraphs[0]
        r2 = p2.add_run(a)
        r2.font.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(4, 120, 87)

    doc.add_page_break()

    # =========================================================================
    # PAGE 3 : LES 4 PILIERS DE L'ÉCOSYSTÈME & MODULE CLIENT
    # =========================================================================
    add_sec_title("3. Les 4 Piliers Fondateurs de l'Écosystème Sellify.me")

    pillars = [
        ("Pilier 1 : La Confiance par l'Escrow Mobile Money", "Moteur transactionnel qui bloque les fonds du client sur un compte séquestre lors de la commande et ne les débloque au vendeur qu'après validation du code secret OTP lors de la remise du colis."),
        ("Pilier 2 : La Logistique Intelligente OSRM & Télémétrie", "Système de routage et de guidage GPS turn-by-turn intégré à l'application. La carte Leaflet suit la géométrie exacte des axes routiers de Douala et Yaoundé avec calcul dynamique d'ETA."),
        ("Pilier 3 : La Suite SaaS Marchande & Smart-Links", "Outils professionnels pour les vendeurs : gestion multi-boutiques, suivi des stocks avec alertes de seuil, grand livre financier avec exports CSV et génération de Smart-Links 1-clic pour les réseaux sociaux."),
        ("Pilier 4 : Le Moteur d'IA Universel Sellify AI 1.2 Flash", "Assistant conversationnel multimodal (texte et voix) déployé sur tous les tableaux de bord pour assister l'acheteur, conseiller le vendeur sur ses prix, guider le livreur et détecter les fraudes."),
    ]

    for p_title_text, p_desc_text in pillars:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(4)
        r_p1 = p_item.add_run(f"• {p_title_text} : ")
        r_p1.font.bold = True
        r_p1.font.size = Pt(9.5)
        r_p1.font.color.rgb = RGBColor(15, 23, 42)
        r_p2 = p_item.add_run(p_desc_text)
        r_p2.font.size = Pt(9)
        r_p2.font.color.rgb = RGBColor(71, 85, 105)

    add_sec_title("4. Spécifications Détaillées des Modules Applicatifs")

    add_subsec_title("4.1 Module Client (Acheteur)")
    doc.add_paragraph(
        "Le module client assure une expérience d'achat moderne, sécurisée et fluide :"
    )
    doc.add_paragraph("• Inscription rapide par téléphone avec validation OTP SMS ou connexion Google OAuth.")
    doc.add_paragraph("• Catalogue interactif multi-boutiques avec moteur de recherche en direct et filtres par zone géographique.")
    doc.add_paragraph("• Panier intelligent calculant automatiquement les remises sur quantité et les frais de livraison par quartier.")
    doc.add_paragraph("• Paiement sécurisé par Escrow Mobile Money (Orange Money / MTN MoMo).")
    doc.add_paragraph("• Suivi de livraison sur carte interactive temps réel avec mise à jour continue de la position du chauffeur.")
    doc.add_paragraph("• Confirmation de réception sécurisée par code secret OTP à 6 chiffres.")
    doc.add_paragraph("• Programme de fidélité récompensant chaque achat (1 point par 100 FCFA dépensés).")

    add_subsec_title("4.2 Module Vendeur (Suite SaaS & Smart-Links)")
    doc.add_paragraph(
        "Le module vendeur fournit une suite complète de gestion commerciale :"
    )
    doc.add_paragraph("• Espace multi-boutiques centralisé permettant de piloter plusieurs points de vente depuis un compte unique.")
    doc.add_paragraph("• Gestion des produits avec déclinaison de variantes (tailles, couleurs, capacités) et suivi d'inventaire.")
    doc.add_paragraph("• Smart-Links de Vente Sociale : Liens de paiement ultra-rapides permettant aux clients WhatsApp/TikTok d'acheter en 1 clic.")
    doc.add_paragraph("• Portefeuille Financier Complet : Séparation claire entre solde disponible et solde sous séquestre, graphiques d'évolution des flux sur 6 mois, demande de retrait instantanée vers Mobile Money.")
    doc.add_paragraph("• Exportation comptable en CSV (Excel) et impression de Relevés de Compte Officiels certifiés (PDF).")
    doc.add_paragraph("• Gestion des codes promo personnalisés et campagnes promotionnelles.")

    doc.add_page_break()

    # =========================================================================
    # PAGE 4 : LIVREUR, SUPERADMIN & CYCLE DE VIE ESCROW
    # =========================================================================
    add_subsec_title("4.3 Module Livreur (Chauffeur & Télémétrie)")
    doc.add_paragraph(
        "Le module livreur optimise le travail des coursiers urbains indépendants :"
    )
    doc.add_paragraph("• Vérification et agrément KYC des livreurs (permis, carte d'identité, immatriculation du véhicule).")
    doc.add_paragraph("• Bascule de disponibilité (En ligne / Hors ligne) et réception intelligente des missions de livraison.")
    doc.add_paragraph("• Guidage GPS Turn-by-Turn intégré directement dans l'application avec calcul d'itinéraire OSRM temps réel.")
    doc.add_paragraph("• Clôture sécurisée avec triple validation : Saisie de l'OTP secret du client + signature tactile sur écran + photo preuve.")
    doc.add_paragraph("• Portefeuille livreur avec rémunération à la course et 100 points de récompense par livraison réussie.")

    add_subsec_title("4.4 Module SuperAdmin (Supervision & Arbitrage)")
    doc.add_paragraph(
        "Le module d'administration générale garantit l'intégrité de la plateforme :"
    )
    doc.add_paragraph("• Tableau de bord exécutif synthétisant le volume global sous séquestre, les commissions et les litiges.")
    doc.add_paragraph("• Centre de validation KYC des marchands et des livreurs partenaires.")
    doc.add_paragraph("• Module d'arbitrage impartial des litiges Escrow (libération forcée vendeur ou remboursement acheteur).")
    doc.add_paragraph("• Journal d'audit immuable (ActivityLog) traçant chaque opération avec adresse IP et horodatage.")

    add_sec_title("5. Mécanique du Séquestre Escrow Mobile Money")

    doc.add_paragraph(
        "Le cœur de la sécurité de Sellify.me repose sur la classe EscrowService.php qui orchestre des transactions atomiques :"
    )

    escrow_steps = [
        ("1. Consignation sous Séquestre", "Dès que le client paie par Mobile Money, les fonds sont crédités sur le pending_balance (solde bloqué) du vendeur. La commande passe à l'état 'escrow_held' et un code secret OTP à 6 chiffres est généré."),
        ("2. Préparation & Dispatch", "Le vendeur prépare la commande et la remet au livreur assigné qui débute l'acheminement avec guidage GPS OSRM."),
        ("3. Validation & Déblocage Automatique", "À la livraison, la saisie du code OTP secret par le livreur ou le clic de confirmation par le client transfère instantanément les fonds de pending_balance vers balance (solde disponible retirable). La commande passe à 'released' et 'delivered'."),
        ("4. Restitution en Cas d'Annulation", "En cas de litige résolu en faveur de l'acheteur ou d'annulation avant départ, EscrowService restitue l'intégralité des fonds à l'acheteur et remet les articles en stock."),
    ]

    for s_num, s_txt in escrow_steps:
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_after = Pt(3)
        r_s1 = p_s.add_run(f"• {s_num} : ")
        r_s1.font.bold = True
        r_s1.font.size = Pt(9)
        r_s1.font.color.rgb = RGBColor(15, 23, 42)
        r_s2 = p_s.add_run(s_txt)
        r_s2.font.size = Pt(8.5)
        r_s2.font.color.rgb = RGBColor(71, 85, 105)

    add_callout(
        "Idempotence & Intégrité Transactionnelle",
        "Toutes les opérations financières s'exécutent au sein de transactions de base de données PostgreSQL isolées (DB::transaction). "
        "Il est mathématiquement impossible qu'une transaction soit débloquée ou remboursée deux fois."
    )

    doc.add_page_break()

    # =========================================================================
    # PAGE 5 : LOGISTIQUE, IA 1.2 FLASH & ARCHITECTURE TECHNIQUE
    # =========================================================================
    add_sec_title("6. Logistique Routière OSRM & Guidage In-App")

    doc.add_paragraph(
        "Contrairement aux solutions traditionnelles qui redirigent vers Google Maps ou Waze, Sellify.me intègre un système "
        "de navigation géographique totalement autonome :"
    )
    doc.add_paragraph("• Moteur OSRM (Open Source Routing Machine) modélisant la géométrie exacte des voiries de Douala et Yaoundé.")
    doc.add_paragraph("• Découpage dynamique de polyligne : Tracé vert pour la portion parcourue et jaune pour le trajet restant.")
    doc.add_paragraph("• Télémétrie en direct avec animation continue de la position du livreur et calcul permanent de l'ETA.")
    doc.add_paragraph("• Interface HUD (Head-Up Display) affichant les instructions de guidage et les boutons d'appel d'urgence.")

    add_sec_title("7. Moteur d'Intelligence Artificielle Sellify AI 1.2 Flash")

    doc.add_paragraph(
        "L'assistant intelligent Sellify AI 1.2 Flash agit comme un copilote universel accessible par texte et par la voix :"
    )
    doc.add_paragraph("• Reconnaissance vocale (Speech-to-Text) et synthèse vocale naturelle (Text-to-Speech).")
    doc.add_paragraph("• Copilote Acheteur : Recherche conversationnelle, localisation instantanée de commande et assistance litige.")
    doc.add_paragraph("• Copilote Vendeur : Recommandation d'optimisation de prix, rédaction de descriptions produits et simulation de marge.")
    doc.add_paragraph("• Copilote Livreur : Dictée vocale des itinéraires, alertes d'embouteillages et synthèse des gains.")
    doc.add_paragraph("• Copilote SuperAdmin : Détection proactive des anomalies de paiement et aide à la décision lors des litiges.")

    add_sec_title("8. Architecture Technique, Sécurité & Performances")

    doc.add_paragraph(
        "La conception logicielle de Sellify.me répond aux exigences les plus strictes du génie logiciel :"
    )

    # Tech Stack Table
    tbl_tech = doc.add_table(rows=6, cols=2)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_tech.autofit = False

    tech_data = [
        ("Backend Framework", "Laravel 11.x (PHP 8.3+) avec architecture Service Layer et Repositories."),
        ("Frontend Réactif", "React 19 + Inertia.js (Single Page Application ultra-fluide sans latence REST)."),
        ("Base de Données", "PostgreSQL 16 avec typage strict, indexation géodésique et contraintes d'intégrité."),
        ("Cartographie & Routage", "Leaflet.js + OpenStreetMap + Serveur de calcul d'itinéraires OSRM haute disponibilité."),
        ("Design System & UI", "Tailwind CSS v4 avec système de cartes Bento, micro-animations et gestion des curseurs."),
        ("Validation & Tests", "Suite automatisée PHPUnit validée à 100% (47 tests unitaires et d'intégration, 255 assertions)."),
    ]

    for idx, (k, v) in enumerate(tech_data):
        c0, c1 = tbl_tech.cell(idx, 0), tbl_tech.cell(idx, 1)
        c0.width, c1.width = Inches(2.2), Inches(4.5)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(8.5)
        r0.font.color.rgb = RGBColor(15, 23, 42)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_page_break()

    # =========================================================================
    # PAGE 6 : BUSINESS MODEL, ROADMAP & CONCLUSION
    # =========================================================================
    add_sec_title("9. Modèle Économique & Monétisation (Business Model)")

    doc.add_paragraph(
        "Sellify.me s'appuie sur trois flux de monétisation complémentaires et pérennes :"
    )

    # Revenue Table
    tbl_rev = doc.add_table(rows=4, cols=3)
    tbl_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_rev.autofit = False

    rev_headers = ["Flux de Revenu", "Tarification & Structure", "Description & Valeur Ajoutée"]
    for i, h in enumerate(rev_headers):
        c = tbl_rev.cell(0, i)
        set_cell_background(c, "0F172A")
        set_cell_margins(c, top=80, bottom=80, left=90, right=90)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    rev_data = [
        ("1. Commission Escrow", "3 % par transaction réussie", "Sécurisation intégrale de la vente et garantie de déblocage."),
        ("2. Abonnements SaaS Marchands", "Starter (Gratuit), Pro (15 000 F/mois), Enterprise (45 000 F/mois)", "Multi-boutiques illimitées, Smart-Links avancés, Copilote IA et comptabilité."),
        ("3. Marge Logistique", "10 % à 15 % par course", "Frais d'intermédiation et d'optimisation de tournée chauffeur."),
    ]

    for row_idx, (s, p, d) in enumerate(rev_data, start=1):
        c0, c1, c2 = tbl_rev.cell(row_idx, 0), tbl_rev.cell(row_idx, 1), tbl_rev.cell(row_idx, 2)
        c0.width, c1.width, c2.width = Inches(2.2), Inches(2.2), Inches(2.4)
        set_cell_background(c0, "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
        set_cell_background(c1, "FFFFFF")
        set_cell_background(c2, "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
        set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
        set_cell_margins(c2, top=60, bottom=60, left=80, right=80)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(s)
        r0.font.bold = True
        r0.font.size = Pt(8)
        r0.font.color.rgb = RGBColor(15, 23, 42)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(p)
        r1.font.bold = True
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor(202, 138, 4)

        p2 = c2.paragraphs[0]
        r2 = p2.add_run(d)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(71, 85, 105)

    add_sec_title("10. Feuille de Route Stratégique (Roadmap)")

    roadmap_steps = [
        ("Phase 1 : 2026 Q3 (Actuel)", "Lancement officiel à Douala et Yaoundé, déploiement complet du séquestre Mobile Money, des Smart-Links de vente sociale et du copilote Sellify AI 1.2 Flash."),
        ("Phase 2 : 2026 Q4", "Déploiement des applications mobiles natives Android et iOS, extension logistique vers les villes secondaires (Bafoussam, Kribi, Garoua) et partenariats bancaires."),
        ("Phase 3 : 2027 Q1 - Q2", "Expansion panafricaine dans la zone CEMAC (Gabon, Congo, Tchad, Guinée Équatoriale), création de hubs logistiques de proximité et interconnexion transfrontalière."),
    ]

    for r_title, r_txt in roadmap_steps:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.space_after = Pt(4)
        r_r1 = p_r.add_run(f"• {r_title} : ")
        r_r1.font.bold = True
        r_r1.font.size = Pt(9)
        r_r1.font.color.rgb = RGBColor(15, 23, 42)
        r_r2 = p_r.add_run(r_txt)
        r_r2.font.size = Pt(8.5)
        r_r2.font.color.rgb = RGBColor(71, 85, 105)

    add_sec_title("11. Conclusion & Relations Institutionnelles")

    doc.add_paragraph(
        "Sellify.me apporte une réponse technologique souveraine, élégante et concrète aux défis majeurs du commerce africain. "
        "En éliminant l'aléa moral du Cash on Delivery grâce au séquestre Escrow et en offrant aux commerçants et livreurs des outils "
        "numériques professionnels, la plateforme s'impose comme le nouveau standard de confiance du commerce numérique panafricain."
    )

    # Sign-off Box
    tbl_sign = doc.add_table(rows=1, cols=1)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_sign = tbl_sign.cell(0, 0)
    set_cell_background(c_sign, "0F172A")
    set_cell_margins(c_sign, top=140, bottom=140, left=200, right=200)
    p_s = c_sign.paragraphs[0]
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_s.add_run("SELLIFY.ME — BÂTIR LA CONFIANCE DU COMMERCE AFRICAIN\n")
    r_s.font.bold = True
    r_s.font.size = Pt(10)
    r_s.font.color.rgb = RGBColor(234, 179, 8)
    r_s2 = p_s.add_run("Contact officiel Direction : direction@sellify.me · contact@sellify.me\nSiège Social : Boulevard de la Liberté, Akwa, Douala — Cameroun\nSite Officiel : https://sellify.me")
    r_s2.font.size = Pt(8.5)
    r_s2.font.color.rgb = RGBColor(203, 213, 225)

    # Output file
    output_docx = "/home/mr-dims-tech/developpement/developpement_laravel/mr_dims/sellify/Dossier_Presentation_Projet_Sellify.docx"
    doc.save(output_docx)
    print(f"Corporate Word Document successfully created at: {output_docx}")

if __name__ == "__main__":
    create_corporate_docx()
